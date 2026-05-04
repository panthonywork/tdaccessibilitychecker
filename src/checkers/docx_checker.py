from src.standards.matterhorn import get_checkpoints_for_filetype


def check_docx(file_bytes: bytes, filename: str) -> list[dict]:
    try:
        from docx import Document
        from io import BytesIO
    except ImportError as e:
        return [_error_result(str(e))]

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        return [_error_result(f"Could not open Word file: {e}")]

    checkpoints = get_checkpoints_for_filetype("docx")
    results = []
    for cp in checkpoints:
        result = _run_checkpoint(cp, doc)
        results.append(result)
    return results


def _run_checkpoint(cp: dict, doc) -> dict:
    cp_id = cp["id"]
    base = {
        "id": cp_id,
        "group_id": cp["group_id"],
        "group_name": cp["group_name"],
        "description": cp["description"],
        "plain_english": cp["plain_english"],
        "fix_hint": cp["fix_hint"],
    }

    try:
        if cp_id == "06-001":
            status = _check_title(doc)
        elif cp_id == "07-001":
            status = _check_language(doc)
        elif cp_id == "13-001":
            status = _check_headings(doc)
        elif cp_id == "13-004":
            status = _check_heading_hierarchy(doc)
        elif cp_id == "17-001":
            status = _check_image_alt_text(doc)
        elif cp_id == "15-001":
            status = _check_lists(doc)
        elif cp_id == "28-002":
            status = _check_link_text(doc)
        elif cp_id == "14-001":
            status = _check_table_headers(doc)
        else:
            status = "na"
    except Exception:
        status = "na"

    return {**base, "status": status}


def _check_title(doc) -> str:
    core_props = doc.core_properties
    title = getattr(core_props, "title", "") or ""
    return "pass" if title.strip() else "fail"


def _check_language(doc) -> str:
    for para in doc.paragraphs:
        if para.runs:
            lang = para.runs[0]._r.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lang")
            if lang is not None:
                return "pass"
    # Check document default language
    try:
        from docx.oxml.ns import qn
        lang_el = doc.element.find(f".//{qn('w:lang')}")
        return "pass" if lang_el is not None else "fail"
    except Exception:
        return "fail"


def _check_headings(doc) -> str:
    heading_styles = {p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")}
    return "pass" if heading_styles else "fail"


def _check_heading_hierarchy(doc) -> str:
    levels = []
    for para in doc.paragraphs:
        name = para.style.name
        if name.startswith("Heading "):
            try:
                level = int(name.split(" ")[-1])
                levels.append(level)
            except ValueError:
                pass

    for i in range(1, len(levels)):
        if levels[i] > levels[i - 1] + 1:
            return "fail"
    return "pass"


def _check_image_alt_text(doc) -> str:
    from docx.oxml.ns import qn
    from lxml import etree

    inline_images = doc.element.findall(f".//{qn('wp:inline')}")
    drawing_images = doc.element.findall(f".//{qn('wp:anchor')}")
    all_images = inline_images + drawing_images

    if not all_images:
        return "na"

    for img in all_images:
        doc_pr = img.find(f"{qn('wp:docPr')}")
        if doc_pr is None:
            return "fail"
        descr = doc_pr.get("descr", "").strip()
        title = doc_pr.get("title", "").strip()
        if not descr and not title:
            return "fail"
    return "pass"


def _check_lists(doc) -> str:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (text.startswith("- ") or text.startswith("• ")):
            style = para.style.name.lower()
            if "list" not in style:
                return "fail"
    return "pass"


def _check_link_text(doc) -> str:
    bad_phrases = {"click here", "here", "read more", "learn more", "more", "link"}
    from docx.oxml.ns import qn
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            for para in doc.paragraphs:
                for run in para.runs:
                    text = run.text.strip().lower()
                    if text in bad_phrases:
                        return "fail"
    return "pass"


def _check_table_headers(doc) -> str:
    from docx.oxml.ns import qn
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = table.rows[0]
        for cell in first_row.cells:
            # Check if header row style is applied
            tr_pr = first_row._tr.find(qn("w:trPr"))
            if tr_pr is None:
                return "fail"
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                return "fail"
    return "pass"


def _error_result(message: str) -> dict:
    return {
        "id": "ERR",
        "group_id": "00",
        "group_name": "Error",
        "description": message,
        "plain_english": message,
        "fix_hint": "Ensure the file is a valid .docx file.",
        "status": "fail",
    }
